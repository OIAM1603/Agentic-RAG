import os
import re
import pandas as pd
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from transformers import AutoTokenizer
from tqdm import tqdm

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain.agents import Tool, initialize_agent, AgentType

from tavily import TavilyClient
import gradio as gr


# ========= CẤU HÌNH HỆ THỐNG =========
os.environ["TAVILY_API_KEY"] = "your-api-key"
os.environ["GOOGLE_API_KEY"] = "your-api-key"

DATA_FOLDER = Path("dataset")
VECTORSTORE_PATH = Path("./dataset/vectorstores")
VECTORSTORE_PATH.mkdir(parents=True, exist_ok=True)

# ========= KHỞI TẠO MÔ HÌNH =========
tokenizer = AutoTokenizer.from_pretrained("intfloat/multilingual-e5-base")
splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
    tokenizer, chunk_size=512, chunk_overlap=64
)
embedding_model = HuggingFaceEmbeddings(model_name="intfloat/multilingual-e5-base")
llm = ChatGoogleGenerativeAI(model="gemini-2.0-flash")
tavily = TavilyClient(api_key=os.environ["TAVILY_API_KEY"])

# ========= HÀM XỬ LÝ DỮ LIỆU ========= 
def clean_text(text: str) -> str:
    text = text.replace("\n", " ").replace("\u2003", " ")
    return re.sub(r"\s+", " ", text).strip()

def load_file(filepath: str) -> list[Document]:
    ext = os.path.splitext(filepath)[1].lower()
    filename = Path(filepath).stem
    base_metadata = {"filename": filename, "filetype": ext, "source_path": str(filepath)}

    if ext == ".txt":
        with open(filepath, encoding="utf-8") as f:
            return [Document(clean_text(f.read()), metadata=base_metadata)]

    elif ext == ".docx":
        doc = DocxDocument(filepath)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        return [Document(clean_text(full_text), metadata=base_metadata)]

    elif ext == ".pdf":
        reader = PdfReader(filepath)
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        return [Document(clean_text(text), metadata=base_metadata)]

    elif ext in [".csv", ".xlsx"]:
        df = pd.read_csv(filepath) if ext == ".csv" else pd.read_excel(filepath, engine="openpyxl")
        df_clean = df.map(lambda x: clean_text(str(x)) if pd.notnull(x) else "")
        return [
            Document(" ".join(str(cell) for cell in row if cell), metadata=base_metadata)
            for _, row in df_clean.iterrows()
        ]

    print(f"[Bỏ qua] Không hỗ trợ định dạng: {filepath}")
    return []

def build_vectorstore():
    all_chunks = []
    filepaths = list(DATA_FOLDER.iterdir())
    print(f"Đang xử lý {len(filepaths)} tệp...")

    for filepath in tqdm(filepaths, desc="Đang xử lý tài liệu", unit="file"):
        if filepath.is_file():
            docs = load_file(str(filepath))
            if docs:
                chunks = splitter.split_documents(docs)
                all_chunks.extend(chunks)

    if all_chunks:
        vectordb = FAISS.from_documents(all_chunks, embedding_model)
        vectordb.save_local(str(VECTORSTORE_PATH))
        print(f"FAISS đã được lưu tại: {VECTORSTORE_PATH}")
    else:
        print("Không có dữ liệu để xây dựng FAISS.")

# ========= CÁC TOOL CHO AGENT =========
def search_knowledge(query: str, source: str = None, filetype: str = None) -> str:
    results = vectordb.similarity_search_with_score(query, k=10)
    filtered = [
        (doc, score)
        for doc, score in results
        if (not source or doc.metadata.get("filename") == source) and
           (not filetype or doc.metadata.get("filetype") == filetype)
    ]
    if not filtered:
        return "Không tìm thấy kết quả phù hợp trong nguồn chỉ định."
    top_results = sorted(filtered, key=lambda x: x[1])[:5]
    return "\n\n---\n\n".join(
        f"[Nguồn: {doc.metadata.get('filename')}]\n{doc.page_content}"
        for doc, _ in top_results
    )

def web_search(query: str, num_results: int = 5) -> str:
    results = tavily.search(query=query, max_results=num_results)["results"]
    if not results:
        return "Không tìm thấy kết quả nào từ internet."
    return "\n\n---\n\n".join(
        f"[URL: {res['url']}]\n{res['content']}" for res in results
    )

# ========= KHỞI TẠO AGENT =========
def create_agent():
    tools = [
        Tool.from_function(
            func=search_knowledge,
            name="search_knowledge",
            description="Tìm kiếm thông tin từ cơ sở dữ liệu nội bộ."
        ),
        Tool.from_function(
            func=web_search,
            name="web_search",
            description="Tìm kiếm thông tin trên internet."
        )
    ]

    system_prompt = PromptTemplate.from_template("""
    Bạn là trợ lý AI thông minh, trả lời bằng tiếng Việt tự nhiên, rõ ràng.
    - Bắt buộc sử dụng hoàn toàn bằng tiếng Việt, không dùng bất kỳ ngôn ngữ nào kháckhác                          
    - Trích dẫn nguồn nếu có (file hoặc URL).
    - Nếu không có dữ liệu, hãy nói rõ không tìm thấy.
    - Giữ lịch sử trò chuyện.
    Lịch sử hội thoại:
    {chat_history}
    Câu hỏi mới:
    {input}
    """)

    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

    return initialize_agent(
        tools=tools,
        llm=llm,
        agent=AgentType.ZERO_SHOT_REACT_DESCRIPTION,
        verbose=True,
        memory=memory,
        system_prompt=system_prompt,
        handle_parsing_errors=True,
    )

# ========= CHAT HANDLER =========
def agent_chat(message, history):
    try:
        print(f"[INPUT]: {message}")
        result = agent.invoke({"input": message})
        return result.get("output", str(result))
    except Exception as e:
        print(f"[ERROR]: {e}")
        return f"Đã xảy ra lỗi: {e}"

# ========= MAIN (CHẠY DEMO GIAO DIỆN) =========
def main():
    global vectordb, agent

    # Nếu vectorstore chưa tồn tại, tạo mới
    if not (VECTORSTORE_PATH / "index.faiss").exists():
        print("Đang xây dựng FAISS...")
        build_vectorstore()

    # Load FAISS
    vectordb = FAISS.load_local(VECTORSTORE_PATH, embedding_model, allow_dangerous_deserialization=True)
    print("FAISS đã được nạp.")

    # Tạo Agent
    agent = create_agent()

    # Giao diện Gradio
    demo = gr.ChatInterface(
        fn=agent_chat,
        title="Trợ lý Agentic RAG",
        description="Đặt câu hỏi bằng tiếng Việt. Agent sẽ tìm câu trả lời từ tài liệu hoặc internet.",
        theme=gr.themes.Default()
    )
    demo.launch()

if __name__ == "__main__":
    main()
